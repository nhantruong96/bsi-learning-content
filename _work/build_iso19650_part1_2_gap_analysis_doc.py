from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Users\nhan9\Documents\Codex\bsi-learning-content")
OUT = ROOT / "_work" / "ISO19650_part1_2_gap_analysis_BIM7192.docx"


def set_styles(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(11)


def add_label(doc: Document, label: str, text: str = ""):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    if text:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    set_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PHÂN TÍCH DECK BIM COORDINATOR & BIM MANAGER THEO ISO 19650-1 & ISO 19650-2")
    r.bold = True
    r.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bản đề xuất chỉnh sửa nội dung slide để người dùng tự cập nhật PPTX")
    r.italic = True

    add_label(doc, "Deck phân tích: ", "02 Slides_BIM7192VNM_v1_Mar 2025.pptx")
    add_label(doc, "Tài liệu đối chiếu: ", "BS EN ISO 19650-1:2018 và BS EN ISO 19650-2:2018 trong thư mục References/ISO 19650.")
    add_label(doc, "Phạm vi: ", "Chỉ đánh giá theo ISO 19650-1 và ISO 19650-2. Các nội dung Part 3, 4, 5, 6 chỉ nên dùng như bối cảnh hoặc ví dụ so sánh nếu cần.")

    doc.add_heading("1. Kết luận chính", level=1)
    add_bullets(doc, [
        "Deck hiện tại đang mô tả BIM Coordinator và BIM Manager như hai chức danh nhân sự tổng quát, nhưng ISO 19650-1 và ISO 19650-2 không chuẩn hóa hai job title này.",
        "ISO 19650 sử dụng mô hình bên tham gia và nhóm: appointing party, lead appointed party, appointed party, project team, delivery team, task team.",
        "ISO 19650 nhấn mạnh information management function. Đây là chức năng/nhiệm vụ được phân bổ cho bên phù hợp, không mặc định là một cá nhân tên BIM Manager hay BIM Coordinator.",
        "Mô tả hiện tại đang trộn lẫn trách nhiệm của client/appointing party, lead appointed party, task team và người thao tác phần mềm. Nếu giữ nguyên, học viên dễ hiểu sai rằng BIM Manager hoặc BIM Coordinator chịu toàn bộ trách nhiệm tiêu chuẩn, BEP, CDE, mô hình, clash, đào tạo và đánh giá.",
        "Cần chỉnh deck theo hướng: giữ BIM Coordinator/BIM Manager như tên chức danh thị trường, nhưng map rõ vào chức năng quản lý thông tin và bối cảnh bổ nhiệm theo ISO 19650.",
    ])

    doc.add_heading("2. Cách hiểu đúng theo ISO 19650-1 và ISO 19650-2", level=1)
    add_table(doc, ["Khái niệm ISO", "Cách hiểu đúng khi đưa vào deck", "Hệ quả khi sửa nội dung"],
              [
                  ["Appointing party", "Bên nhận thông tin từ lead appointed party; trong dự án thường tương ứng client/chủ đầu tư/bên đặt hàng.", "Không gán toàn bộ trách nhiệm yêu cầu thông tin cho BIM Manager của nhà thầu nếu người đó không đại diện appointing party."],
                  ["Lead appointed party", "Bên được appointing party chỉ định trực tiếp và dẫn dắt một delivery team.", "Các nội dung pre-appointment BEP, capability/capacity, mobilization plan, risk register, MIDP thường thuộc lead appointed party trong quy trình ISO 19650-2."],
                  ["Appointed party", "Bên cung cấp thông tin, hàng hóa hoặc dịch vụ; có thể là task team trong delivery team.", "Không mô tả mọi appointed party như BIM Coordinator; họ có thể là bộ môn hoặc nhà thầu phụ tạo thông tin."],
                  ["Delivery team / task team", "Delivery team gồm lead appointed party và appointed parties; task team là nhóm thực hiện nhiệm vụ cụ thể.", "Việc tạo thông tin/mô hình thuộc task team; BIM Coordinator có thể điều phối, kiểm tra, tổng hợp, nhưng không mặc định tự tạo mọi mô hình."],
                  ["Information management function", "Chức năng quản lý thông tin được phân bổ cho cá nhân/tổ chức có phạm vi, trách nhiệm và quyền hạn rõ.", "Nên thay logic job title bằng function/responsibility matrix. Có thể nói BIM Manager/Coordinator là cách doanh nghiệp đặt tên cho một phần chức năng này."],
                  ["Responsibility matrix", "Ma trận phân bổ trách nhiệm cho function và information deliverables.", "Cần thay danh sách mô tả công việc chung bằng ma trận trách nhiệm theo appointment và theo bên tham gia."],
                  ["EIR / BEP / TIDP / MIDP", "EIR do appointing party và lead appointed party thiết lập ở các cấp; BEP giải thích cách delivery team thực hiện quản lý thông tin; TIDP/MIDP lập kế hoạch giao nộp information containers.", "Không mô tả BEP như tài liệu riêng của BIM Manager; phải gắn với delivery team và appointment."],
                  ["CDE", "Nguồn thông tin thống nhất có workflow quản lý information container, trạng thái, metadata, review/authorize/accept.", "Không mô tả CDE chỉ như nền tảng phần mềm. Cần nhấn mạnh quy trình và trạng thái thông tin."],
                  ["Capability and capacity", "Đánh giá khả năng và năng lực đáp ứng yêu cầu thông tin, quy trình, công cụ, nhân sự, đào tạo và nguồn lực.", "Không đánh giá chức danh chủ yếu bằng chứng chỉ, số năm kinh nghiệm hoặc thành thạo phần mềm."],
              ])

    doc.add_heading("3. Các điểm chưa phù hợp trong deck hiện tại", level=1)
    gap_rows = [
        ["Slide 6", "Mục đích khóa học nói 'cung cấp bằng chứng đánh giá năng lực phù hợp' cho người phối hợp thiết lập mô hình.", "ISO 19650-2 nói về capability/capacity và information management function theo appointment, không phải chứng minh năng lực cho hai chức danh cố định.", "Sửa thành: khóa học cung cấp khung tham khảo để xác định, phân bổ và đánh giá năng lực thực hiện information management function trong delivery phase."],
        ["Slide 7", "Mục tiêu học tập thiếu mô hình parties/teams và information management function.", "Đây là nền tảng của ISO 19650-1/2. Nếu chỉ nói phần mềm, BEP, mô tả job sẽ làm lệch trọng tâm.", "Bổ sung mục tiêu: phân biệt appointing party, lead appointed party, appointed party, delivery team, task team và cách function được phân bổ."],
        ["Slide 9-13", "Phần bối cảnh và định nghĩa BIM còn thiên về 'mô hình' và 'triển khai BIM' chung.", "ISO 19650 định vị BIM là shared digital representation để hỗ trợ quyết định, nằm trong information management.", "Sửa phần định nghĩa: BIM phục vụ quản lý thông tin theo yêu cầu thông tin, mô hình thông tin và quyết định của appointing party."],
        ["Slide 12", "Đưa nhiều Part của ISO 19650 vào bản đồ nhưng khóa đang yêu cầu Part 1/2.", "Có thể gây hiểu là khóa bao phủ Part 3-6.", "Giữ Part 3-6 như bối cảnh, nhưng đánh dấu rõ phạm vi chính là Part 1 concepts/principles và Part 2 delivery phase."],
        ["Slide 18-24", "BEP/CDE được trình bày tương đối đúng nhưng thiếu bước Assessment and need của appointing party.", "ISO 19650-2 bắt đầu từ appointing party: information requirements, milestones, information standard, production methods/procedures, reference info/shared resources, project CDE, information protocol.", "Thêm 1 slide hoặc khối nội dung trước BEP: trách nhiệm chuẩn bị của appointing party trước invitation to tender."],
        ["Slide 19", "BEP trước chỉ định có thể bị hiểu là trách nhiệm của BIM Manager/Coordinator.", "Theo ISO 19650-2, prospective lead appointed party lập pre-appointment BEP để đưa vào tender response.", "Ghi rõ subject: prospective lead appointed party/delivery team. BIM Manager/Coordinator chỉ là người được giao hỗ trợ nếu appointment quy định."],
        ["Slide 20-21", "Kế hoạch huy động, risk register, appointment, TIDP/MIDP chưa được kết nối đủ với lead appointed party/task team.", "ISO 19650-2 phân rõ lead appointed party lập mobilization plan, risk register, detailed responsibility matrix, lead appointed party EIR, TIDP và MIDP.", "Bổ sung sơ đồ chuỗi: pre-BEP -> capability/capacity -> mobilization plan/risk register -> appointment -> responsibility matrix -> TIDP/MIDP."],
        ["Slide 22", "CDE dễ bị hiểu là platform.", "ISO coi CDE là nguồn thông tin thống nhất và workflow quản lý information containers.", "Thêm workflow states: WIP, Shared, Published, Archive; metadata status/suitability, revision, classification; review/authorize/accept."],
        ["Slide 25-36", "Mô tả job hiện tại gom quá nhiều trách nhiệm vào BIM Coordinator/BIM Manager, bao gồm phát triển tiêu chuẩn cho tất cả dự án, tạo/bảo trì mô hình, đào tạo, tuân thủ, quản lý dữ liệu.", "ISO phân bổ trách nhiệm theo parties và functions. Task team tạo thông tin; lead appointed party tổng hợp, review/authorize; appointing party accept.", "Thay bằng ba nhóm vai trò ISO-aligned: Appointing party information management function; Lead appointed party information management function; Task team/delivery coordination function."],
        ["Slide 27-30", "BIM Coordinator bị mô tả như vừa điều phối, vừa tạo mô hình, vừa quản lý tiêu chuẩn, vừa đào tạo.", "Theo ISO, tạo thông tin thuộc task team; điều phối và review phải nằm trong responsibility matrix.", "Sửa Coordinator thành vai trò hỗ trợ task/delivery information coordination: kiểm tra information containers, theo dõi issue, phối hợp clash, hỗ trợ CDE metadata, tổng hợp TIDP input."],
        ["Slide 32-36", "BIM Manager bị mô tả như người sở hữu toàn bộ tiêu chuẩn, công cụ, BEP, mô hình và đào tạo.", "ISO không nêu BIM Manager; nếu dùng title này phải chỉ rõ đang đại diện appointing party hay lead appointed party.", "Tách thành BIM Manager phía appointing party và BIM Manager phía lead appointed party, hoặc đổi tên thành Information Management Lead tùy ngữ cảnh."],
        ["Slide 38-41", "Kỹ năng nhấn mạnh phần mềm Revit/AutoCAD/Navisworks.", "ISO yêu cầu khả năng đáp ứng information requirements và quy trình, phần mềm chỉ là phương tiện nếu appointment yêu cầu.", "Bổ sung kỹ năng bắt buộc theo ISO: đọc EIR/PIR, level of information need, information standard, methods/procedures, CDE workflow, responsibility matrix, TIDP/MIDP, QA/review/authorization."],
        ["Slide 44-47", "Đánh giá trình độ đang nghiêng về bằng cấp, chứng chỉ, phỏng vấn kỹ thuật và kiểm tra phần mềm.", "ISO 19650-1/2 nói capability/capacity review của delivery team/task team và competency của cá nhân thực hiện information management function.", "Chuyển thành capability & capacity assessment: tổ chức, quy trình, công cụ, nhân sự, kinh nghiệm collaboration, khả năng giao nộp thông tin, training/mobilization, no conflict of interest."],
    ]
    add_table(doc, ["Slide/cụm slide", "Vấn đề hiện tại", "Vì sao chưa phù hợp ISO 19650-1/2", "Đề xuất chỉnh"], gap_rows)

    doc.add_heading("4. Nếu áp dụng ISO 19650-1 & 2 thì mô tả vai trò nên như thế nào?", level=1)
    doc.add_paragraph("Khuyến nghị không bỏ tên BIM Coordinator/BIM Manager nếu thị trường hoặc doanh nghiệp vẫn dùng. Tuy nhiên, cần ghi rõ đây là tên chức danh nội bộ/thị trường được map vào information management function, không phải thuật ngữ vai trò chuẩn hóa của ISO 19650.")

    role_rows = [
        ["Appointing Party Information Management Function", "Đại diện bên đặt hàng/chủ đầu tư về quản lý thông tin.", "Thiết lập PIR, mốc giao nộp thông tin, information standard, production methods/procedures, reference information/shared resources, project CDE, information protocol; lập EIR và tiêu chí đánh giá tender; review/accept information model.", "Có thể là Client Information Manager, Employer's Information Manager, hoặc BIM Manager phía chủ đầu tư nếu tổ chức dùng tên này."],
        ["Lead Appointed Party Information Management Function", "Đại diện bên được chỉ định chính/delivery team.", "Lập pre-appointment BEP; đánh giá capability/capacity của task teams; lập mobilization plan và risk register; sau appointment xác nhận BEP, detailed responsibility matrix, lead appointed party EIR, TIDP/MIDP; review/authorize information model trước khi submit.", "Có thể là BIM Manager/Digital Delivery Manager của tư vấn chính, tổng thầu hoặc nhà thầu chính."],
        ["Task Team Information Management / Coordination Function", "Chức năng phối hợp thông tin ở cấp bộ môn/task team.", "Lập TIDP input; tạo/giao nộp information containers theo methods/procedures; kiểm tra chất lượng trước khi chia sẻ; xử lý coordination issues; cập nhật metadata/status; submit thông tin cho lead appointed party.", "Có thể map với BIM Coordinator, Model Coordinator, Discipline BIM Lead."],
        ["Information Author / Model Author", "Người hoặc nhóm tạo thông tin/mô hình.", "Tạo thông tin theo yêu cầu, level of information need, naming convention, CDE workflow, TIDP và tiêu chuẩn dự án.", "Không nên gộp mặc định vào BIM Coordinator nếu tổ chức có modelers/designers riêng."],
    ]
    add_table(doc, ["Vai trò ISO-aligned nên dùng", "Bản chất", "Trách nhiệm chính theo ISO 19650-1/2", "Có thể map với chức danh hiện hữu"], role_rows)

    doc.add_heading("5. Đề xuất cấu trúc slide sửa lại", level=1)
    structure_rows = [
        ["Phần 1", "Tổng quan ISO 19650-1/2", "Bổ sung mô hình parties/teams và information management function ngay sau phần định nghĩa BIM."],
        ["Phần 2", "Công cụ và nền tảng", "Đổi thông điệp: công cụ hỗ trợ CDE, information production, QA, coordination; không là yêu cầu năng lực mặc định."],
        ["Phần 3", "Quy trình ISO 19650-2 delivery phase", "Trình bày theo quy trình 5.1 đến 5.8: Assessment and need, Invitation to tender, Tender response, Appointment, Mobilization, Collaborative production, Delivery, Close-out."],
        ["Phần 4", "Vai trò và trách nhiệm", "Thay job descriptions generic bằng responsibility map: appointing party IMF, lead appointed party IMF, task team coordination, information authors."],
        ["Phần 5", "Kỹ năng", "Sắp xếp kỹ năng theo function: yêu cầu thông tin, BEP, CDE workflow, responsibility matrix, TIDP/MIDP, QA/review, communication, coordination."],
        ["Phần 6", "Đánh giá năng lực", "Đổi từ đánh giá cá nhân theo job title sang capability/capacity assessment theo appointment và function."],
    ]
    add_table(doc, ["Phần deck", "Tên đề xuất", "Điểm chỉnh chính"], structure_rows)

    doc.add_heading("6. Nội dung thay thế gợi ý cho các slide quan trọng", level=1)
    replacement_rows = [
        ["Slide 6 - Mục đích khóa học", "Khóa học giúp học viên hiểu cách ISO 19650-1 và ISO 19650-2 tổ chức chức năng quản lý thông tin trong giai đoạn triển khai dự án, từ đó map các chức danh BIM Coordinator/BIM Manager hiện hữu vào trách nhiệm, quyền hạn và sản phẩm thông tin phù hợp."],
        ["Slide 7 - Mục tiêu học tập", "Sau khóa học, học viên có thể: phân biệt parties/teams theo ISO; giải thích information management function; mô tả luồng EIR-BEP-TIDP-MIDP-CDE; phân bổ trách nhiệm bằng responsibility matrix; đánh giá capability/capacity theo yêu cầu appointment."],
        ["Slide 13 - BIM là gì", "BIM theo ISO nên được trình bày là việc sử dụng biểu diễn số dùng chung của tài sản để hỗ trợ quá trình thiết kế, xây dựng, vận hành và tạo cơ sở đáng tin cậy cho quyết định. Trọng tâm là thông tin dùng chung và quyết định, không chỉ mô hình 3D."],
        ["Slide mới sau Slide 13", "Thêm slide 'Mô hình parties và teams theo ISO 19650': appointing party, lead appointed party, appointed party, project team, delivery team, task team; luồng yêu cầu thông tin và trao đổi thông tin giữa các bên."],
        ["Slide mới trước Slide 19", "Thêm slide 'Trách nhiệm của appointing party trước tender': PIR, information delivery milestones, project information standard, production methods/procedures, reference information/shared resources, project CDE, information protocol."],
        ["Slide 19", "Đổi tiêu đề thành 'Pre-appointment BEP của prospective lead appointed party'. Nội dung nêu delivery team's proposed approach, information delivery strategy, federation strategy, high-level responsibility matrix, proposed amendments, software/hardware/IT schedule."],
        ["Slide 22", "Đổi nội dung thành 'CDE workflow và information container': WIP, Shared, Published, Archive; status/suitability, revision, classification; review/authorize/accept; audit trail."],
        ["Slide 27-30", "Đổi 'BIM Coordinator - mô tả công việc' thành 'Task team / Delivery information coordination function'. Nội dung tập trung vào TIDP input, information container QA, CDE metadata, coordination issue tracking, model federation support, submit to lead appointed party."],
        ["Slide 32-36", "Đổi 'BIM Manager - mô tả công việc' thành hai ngữ cảnh: 'Information management function - appointing party side' và 'Information management function - lead appointed party side'. Không trộn client acceptance và delivery authorization trong cùng một mô tả."],
        ["Slide 38-41", "Thay checklist phần mềm bằng competency groups: ISO concepts, information requirements, responsibility matrix, BEP, CDE workflow, information delivery planning, QA/review/authorization, communication and coordination."],
        ["Slide 44-47", "Đổi tiêu đề thành 'Capability and capacity assessment'. Đánh giá theo tổ chức/function: kinh nghiệm collaborative working, quy trình, công cụ, nhân sự, training/mobilization, khả năng đáp ứng EIR, TIDP/MIDP, risk register, authority và conflict of interest."],
    ]
    add_table(doc, ["Slide", "Nội dung thay thế gợi ý"], replacement_rows)

    doc.add_heading("7. Phụ lục: hành động chỉnh sửa theo từng slide", level=1)
    slide_action_rows = [
        ["Slide 1-5", "Giữ", "Phần mở đầu, giới thiệu và cấu trúc khóa học không xung đột ISO. Chỉ cần thêm một câu phạm vi: khóa tập trung vào mapping chức danh vào function theo ISO 19650-1/2."],
        ["Slide 6", "Sửa mạnh", "Đổi mục đích từ 'cung cấp bằng chứng đánh giá năng lực' sang 'cung cấp khung tham khảo để xác định và phân bổ information management function'."],
        ["Slide 7", "Sửa", "Bổ sung mục tiêu học tập về parties/teams, information management function, responsibility matrix, EIR-BEP-TIDP-MIDP-CDE."],
        ["Slide 8", "Giữ + mở rộng", "Khi mở Phần 1, nêu rõ trọng tâm ISO 19650-1 là concepts/principles và cách tổ chức thông tin, không chỉ bối cảnh BIM chung."],
        ["Slide 9", "Sửa nhẹ", "Giữ bối cảnh toàn cầu, nhưng tránh nói khóa 'Cơ bản về BIM' nếu deck là BIM Coordinator & Manager. Nêu ISO 19650 là khung quản lý thông tin."],
        ["Slide 10", "Giữ có điều kiện", "Dùng UK/VN như bối cảnh. Nếu phát hành chính thức, cần kiểm tra lại cập nhật pháp lý Việt Nam và không biến slide này thành nội dung pháp lý chi tiết."],
        ["Slide 11", "Giữ + nối ISO", "Giữ lý do dùng tiêu chuẩn, bổ sung rằng tiêu chuẩn giúp thống nhất yêu cầu thông tin, trách nhiệm và quy trình giao nộp."],
        ["Slide 12", "Sửa phạm vi", "Giữ bản đồ tiêu chuẩn nhưng tô rõ phạm vi chính của khóa là ISO 19650-1/2. Part 3-6 chỉ là bối cảnh mở rộng."],
        ["Slide 13", "Sửa định nghĩa", "Thay mô tả BIM thiên về mô hình bằng định nghĩa định hướng ISO: shared digital representation hỗ trợ thiết kế, xây dựng, vận hành và quyết định."],
        ["Slide mới sau 13", "Thêm", "Thêm sơ đồ ISO parties/teams: appointing party, lead appointed party, appointed party, project team, delivery team, task team."],
        ["Slide 14", "Sửa nhẹ", "Nêu rõ phần mềm/công cụ/nền tảng chỉ là phương tiện hỗ trợ information management process."],
        ["Slide 15", "Sửa", "Chuyển 'mục đích BIM' thành vòng đời thông tin và quyết định; liên hệ PIM/AIM ở mức khái niệm nếu cần, không đi sâu Part 3."],
        ["Slide 16", "Giữ như ví dụ", "Dùng BIM Collaboration Pro/Insight như ví dụ platform; tránh tạo cảm giác công cụ này là yêu cầu bắt buộc theo ISO."],
        ["Slide 17", "Sửa hoạt động", "Yêu cầu học viên phân loại công cụ theo quy trình ISO: authoring, CDE, QA, coordination, review/authorize/accept, reporting."],
        ["Slide mới trước 19", "Thêm", "Thêm trách nhiệm appointing party trước tender: PIR, milestones, information standard, production methods/procedures, reference info/shared resources, project CDE, information protocol."],
        ["Slide 18", "Sửa mở phần", "Nêu BEP chỉ là một phần của quy trình ISO 19650-2; không bắt đầu bằng BEP nếu chưa nói assessment and need/invitation to tender."],
        ["Slide 19", "Sửa tiêu đề/chủ thể", "Ghi rõ pre-appointment BEP do prospective lead appointed party lập trong tender response."],
        ["Slide 20", "Sửa nhẹ", "Ghi rõ mobilization plan và information delivery risk register thuộc prospective lead appointed party/delivery team."],
        ["Slide 21", "Mở rộng", "Bổ sung detailed responsibility matrix, lead appointed party EIR, TIDP, MIDP và appointment documents."],
        ["Slide 22", "Sửa mạnh", "Mô tả CDE workflow theo information container: WIP, Shared, Published, Archive; status/suitability, revision, classification; audit trail."],
        ["Slide 23", "Giữ như bối cảnh", "IFC/COBie là bối cảnh open information exchange. Không đi sâu vì nội dung exchange chi tiết thuộc Part 4."],
        ["Slide 24", "Sửa hoạt động", "Hoạt động về BEP nên hỏi: BEP trả lời EIR thế nào, ai chịu trách nhiệm, TIDP/MIDP liên hệ ra sao, CDE workflow nào được áp dụng."],
        ["Slide 25", "Sửa mạnh", "Đổi phần 'công việc điển hình của chức danh' thành 'information management functions và trách nhiệm theo appointment'."],
        ["Slide 26", "Sửa", "Kỹ năng và kiến thức nên gắn với function: information requirements, CDE, responsibility matrix, BEP, TIDP/MIDP, QA/review."],
        ["Slide 27-30", "Sửa mạnh", "BIM Coordinator nên được mô tả như task team/delivery information coordination function: hỗ trợ TIDP, information containers, CDE metadata, issue tracking, QA và coordination."],
        ["Slide 31", "Sửa hoạt động", "So sánh vai trò hiện tại của công ty với ISO functions, không chỉ so sánh hai job title."],
        ["Slide 32-36", "Sửa mạnh", "Tách BIM Manager thành hai ngữ cảnh: appointing party IMF và lead appointed party IMF. Không gộp acceptance của client và authorization của delivery team."],
        ["Slide 37", "Sửa mở phần", "Giới thiệu competency framework theo ISO function, không theo job description chung."],
        ["Slide 38-39", "Sửa", "BIM Coordinator skills: thêm CDE workflow, information container, naming/status, TIDP, QA check, issue coordination; giảm nhấn mạnh phần mềm cụ thể."],
        ["Slide 40-41", "Sửa", "BIM Manager skills: thêm EIR/BEP/MIDP/responsibility matrix, capability/capacity, mobilization, authorization/acceptance workflow, governance."],
        ["Slide 42", "Sửa hoạt động", "Clash detection nên đặt trong information model review/coordination issue workflow, không coi là toàn bộ quản lý BIM."],
        ["Slide 43", "Sửa tiêu đề", "Đổi thành 'Capability and capacity assessment theo ISO 19650-1/2'."],
        ["Slide 44-45", "Sửa mạnh", "Đánh giá Coordinator theo capability/capacity của task/team function, không chỉ chứng chỉ, số năm, thao tác phần mềm."],
        ["Slide 46-47", "Sửa mạnh", "Đánh giá Manager theo function side: appointing party hoặc lead appointed party; bổ sung authority, scope, conflict of interest, ability to manage information process."],
        ["Slide 48", "Sửa hoạt động", "Yêu cầu học viên lập mini responsibility matrix và capability evidence cho từng function."],
        ["Slide 49-51", "Giữ + kết luận ISO", "Kết thúc bằng thông điệp: chức danh có thể khác nhau, nhưng ISO yêu cầu rõ function, responsibility, authority và information deliverables."],
    ]
    add_table(doc, ["Slide", "Hành động", "Nội dung chỉnh cụ thể"], slide_action_rows)

    doc.add_heading("8. Ma trận trách nhiệm đề xuất để đưa vào deck", level=1)
    matrix_rows = [
        ["Xác định PIR, mốc quyết định, yêu cầu thông tin dự án", "Accountable", "Consulted", "Informed/Consulted", "Không gán cho BIM Manager phía nhà thầu nếu không được appointing party giao."],
        ["Lập EIR của appointing party", "Accountable", "Consulted", "Informed", "Có thể có tư vấn hỗ trợ nhưng appointing party giữ trách nhiệm yêu cầu."],
        ["Thiết lập project information standard và production methods/procedures", "Accountable", "Consulted/Apply", "Apply", "Lead appointed party có thể đề xuất amendment trong BEP."],
        ["Thiết lập project CDE và information protocol", "Accountable", "Consulted/Use", "Use", "Có thể thuê bên thứ ba host/manage CDE, nhưng cần scope rõ."],
        ["Lập pre-appointment BEP", "Evaluate", "Accountable", "Input", "Thuộc tender response của prospective lead appointed party."],
        ["Đánh giá task team capability/capacity", "Review criteria", "Accountable", "Provide evidence", "Không chỉ đánh giá cá nhân; phải đánh giá năng lực và nguồn lực nhóm."],
        ["Lập mobilization plan và risk register", "Review/Evaluate", "Accountable", "Input", "Risk register tập trung vào rủi ro giao nộp thông tin."],
        ["Xác nhận BEP sau appointment", "Agree", "Accountable", "Input/Comply", "BEP phải được quản lý qua change control nếu đưa vào appointment documents."],
        ["Lập detailed responsibility matrix", "Review", "Accountable", "Input", "Nên là công cụ chính thay cho job description chung chung."],
        ["Lập TIDP", "Informed", "Coordinate/Aggregate", "Accountable for own tasks", "Task team lập TIDP cho phần việc cụ thể."],
        ["Lập MIDP", "Review", "Accountable", "Input via TIDP", "Lead appointed party aggregate TIDPs."],
        ["Generate information / model", "Informed", "Coordinate", "Accountable", "Thông tin do task team tạo theo methods/procedures."],
        ["QA check và approve for sharing", "Informed", "Coordinate/Review", "Accountable first-line", "Đây là collaborative production process."],
        ["Authorize information model trước khi gửi appointing party", "Informed", "Accountable", "Submit to LAP", "Lead appointed party review/authorize."],
        ["Accept information model", "Accountable", "Submit/Respond", "Support correction", "Appointing party review/accept hoặc reject."],
    ]
    add_table(doc, ["Hoạt động", "Appointing party IMF", "Lead appointed party IMF", "Task team / Coordinator", "Ghi chú"], matrix_rows)

    doc.add_heading("9. Checklist chỉnh PPTX", level=1)
    add_bullets(doc, [
        "Thay mọi câu khẳng định 'BIM Manager/BIM Coordinator chịu trách nhiệm...' bằng 'trong appointment cụ thể, chức năng này có thể được giao cho...'.",
        "Thêm slide giải thích ISO 19650 không chuẩn hóa job title BIM Manager/BIM Coordinator; doanh nghiệp có thể dùng title này để đặt tên cho function.",
        "Tách rõ phía appointing party và phía lead appointed party khi mô tả BIM Manager.",
        "Tách rõ người tạo thông tin/model khỏi người điều phối/quản lý thông tin.",
        "Đưa responsibility matrix vào trung tâm của phần vai trò.",
        "Đổi đánh giá trình độ thành capability and capacity assessment; phần mềm chỉ là một bằng chứng khi EIR/appointment yêu cầu.",
        "Khi nhắc ISO 19650-3/4/5/6, ghi rõ đó là bối cảnh mở rộng, không thuộc phạm vi chính của khóa này nếu đang chỉnh theo Part 1/2.",
    ])

    doc.add_heading("10. Gợi ý thông điệp kết luận cho deck", level=1)
    doc.add_paragraph(
        "Theo ISO 19650-1 và ISO 19650-2, trọng tâm không phải là đặt một người tên BIM Manager hay BIM Coordinator chịu toàn bộ trách nhiệm BIM. "
        "Trọng tâm là xác định đúng bên tham gia, yêu cầu thông tin, chức năng quản lý thông tin, quyền hạn, ma trận trách nhiệm và kế hoạch giao nộp thông tin. "
        "Do đó, nếu doanh nghiệp dùng chức danh BIM Manager hoặc BIM Coordinator, cần map các chức danh này vào appointment, delivery team, task team và responsibility matrix cụ thể."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
