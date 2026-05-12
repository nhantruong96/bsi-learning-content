from __future__ import annotations

import argparse
from pathlib import Path
from shutil import copyfile

from deep_translator import GoogleTranslator
from docx import Document


RAW_EXACT_TRANSLATIONS = {
    "References": "Tài liệu tham khảo",
    "Table of Contents:": "Mục lục:",
    "Activity 1: What are the physical and cyber security risks for a city centre HQ building?\t2": "Hoạt động 1: Rủi ro an ninh vật lý và an ninh mạng đối với tòa nhà trụ sở chính ở trung tâm thành phố là gì?\t2",
    "Activity 2: What is a sensitive asset?\t3": "Hoạt động 2: Tài sản nhạy cảm là gì?\t3",
    "Activity 3: Practical implementation of security mindedness\t4": "Hoạt động 3: Triển khai tư duy an ninh trong thực tế\t4",
    "Activity 4: Responses to a security breach\t5": "Hoạt động 4: Ứng phó với vi phạm an ninh\t5",
    "Activity 5: Action planning for BIM security mindedness\t6": "Hoạt động 5: Lập kế hoạch hành động cho tư duy an ninh của BIM\t6",
    "Activity 1: What are the physical and cyber security risks for a city centre HQ building?": "Hoạt động 1: Rủi ro an ninh vật lý và an ninh mạng đối với tòa nhà trụ sở chính ở trung tâm thành phố là gì?",
    "Activity 2: What is a sensitive asset?": "Hoạt động 2: Tài sản nhạy cảm là gì?",
    "A useful first step is to identify what is a sensitive asset? Possible answers are:": "Bước hữu ích đầu tiên là xác định thế nào là một tài sản nhạy cảm. Một số câu trả lời khả dĩ gồm:",
    "Government Designated – Government ministry buildings": "Tài sản do chính phủ chỉ định - trụ sở các bộ/ngành của chính phủ",
    "Critical national infrastructure (CPNI): International airport": "Hạ tầng quốc gia trọng yếu (CPNI): sân bay quốc tế",
    "Defence, law enfacement, national security or diplomatic: Aldershot Barracks": "Quốc phòng, thực thi pháp luật, an ninh quốc gia hoặc ngoại giao: Doanh trại Aldershot",
    "Landmark or nationally significant site or crowded place: British Museum, Trafalgar Square": "Công trình biểu tượng, địa điểm quan trọng cấp quốc gia hoặc nơi tập trung đông người: Bảo tàng Anh, Quảng trường Trafalgar",
    "Hosting events of security significance: The Mall for state visits": "Địa điểm tổ chức các sự kiện có ý nghĩa về an ninh: The Mall cho các chuyến thăm cấp nhà nước",
    "Significant volume of valuable contents: Bank of England": "Nơi chứa khối lượng lớn tài sản có giá trị: Ngân hàng Anh",
    "How can information be released inadvertently? Some examples are:": "Thông tin có thể bị lộ ngoài ý muốn theo những cách nào? Một số ví dụ gồm:",
    "Publicity (new facilities are frequently documented in press releases and trade press)": "Hoạt động truyền thông (các cơ sở mới thường được giới thiệu trong thông cáo báo chí và báo chí chuyên ngành)",
    "Background detail in photos (corporate brochures, articles)": "Chi tiết hậu cảnh trong ảnh (tài liệu giới thiệu doanh nghiệp, bài viết)",
    "Too large shots of complete BIM models, (front covers, talks)": "Hình chụp quá rộng của toàn bộ mô hình BIM (bìa tài liệu, bài thuyết trình)",
    "Supplier testimonials": "Lời chứng thực của nhà cung cấp",
    "Sub-contractors’ information": "Thông tin từ các nhà thầu phụ",
    "Example: A speculative office building in North London was developed using BIM the 1980’s. It used one of the first general purpose BIM authoring tools, the first 2D/3D integration, the first integrated animated walk-through, the first whole building energy simulations, and complete 3D brick detailing. It was pre-let to a bank. No further mention of the building ever appeared in the press or history of BIM in the UK.": "Ví dụ: Một tòa nhà văn phòng được phát triển theo hình thức đầu cơ tại Bắc Luân Đôn đã được triển khai bằng BIM từ những năm 1980. Công trình này sử dụng một trong những công cụ tạo lập BIM đa dụng đầu tiên, hình thức tích hợp 2D/3D đầu tiên, mô phỏng tham quan ảo tích hợp đầu tiên, các mô phỏng năng lượng cho toàn bộ tòa nhà đầu tiên và cả chi tiết xây gạch 3D đầy đủ. Tòa nhà đã được một ngân hàng thuê trước. Sau đó công trình hầu như không còn được nhắc đến trên báo chí hoặc trong lịch sử BIM tại Vương quốc Anh.",
    "Activity 3: Practical implementation of security mindedness": "Hoạt động 3: Triển khai tư duy an ninh trong thực tế",
    "The security management plan records the organizational policies and processes put in place to manage asset and information security considerations in relation to a particular asset or project.": "Kế hoạch quản lý an ninh ghi nhận các chính sách và quy trình mà tổ chức thiết lập để quản lý các cân nhắc về an ninh tài sản và thông tin đối với một tài sản hoặc dự án cụ thể.",
    "The example solution below is related to a project to build a new terminal building at an international airport. The example clearly does not include details of all potential security considerations but indicates the types of considerations that would be included in the security management plan.": "Ví dụ lời giải dưới đây liên quan đến một dự án xây dựng nhà ga mới tại một sân bay quốc tế. Ví dụ này rõ ràng không bao quát chi tiết toàn bộ các cân nhắc an ninh có thể phát sinh, nhưng cho thấy những loại nội dung cần được đưa vào kế hoạch quản lý an ninh.",
    "Personnel:": "Nhân sự:",
    "What screening and vetting standards will be applied to personnel working on the project? Enhanced level for client-side staff, project directors/team leaders on supply-side. Base level for all other staff.": "Những tiêu chuẩn sàng lọc và thẩm tra nào sẽ được áp dụng cho nhân sự làm việc trong dự án? Mức tăng cường áp dụng cho nhân sự phía khách hàng và giám đốc dự án/trưởng nhóm ở phía nhà cung cấp. Mức cơ bản áp dụng cho toàn bộ nhân sự còn lại.",
    "Mandatory training/induction requirements for all personnel joining the project. All this training is to be controlled by the client security manager.": "Yêu cầu đào tạo/hội nhập bắt buộc đối với mọi nhân sự tham gia dự án. Toàn bộ hoạt động đào tạo này do người quản lý an ninh của khách hàng kiểm soát.",
    "Mandatory de-briefing requirements for all personnel leaving the project or those still in post at project completion. All de-briefing to be controlled by client security manager.": "Yêu cầu rà soát khi kết thúc nhiệm vụ, áp dụng bắt buộc cho mọi nhân sự rời dự án hoặc vẫn còn tại vị trí khi dự án hoàn tất. Toàn bộ hoạt động này do người quản lý an ninh của khách hàng kiểm soát.",
    "The processes for granting and removing CDE access to/from individuals working on the project. Notice to all prospective and appointed suppliers of the time taken to do this.": "Quy trình cấp và thu hồi quyền truy cập CDE cho các cá nhân làm việc trong dự án. Cần thông báo cho tất cả nhà cung cấp dự kiến tham gia và các nhà cung cấp đã được giao nhiệm vụ về thời gian cần thiết để thực hiện việc này.",
    "Physical:": "Vật lý:",
    "How will access control to the work site be maintained and how will the site be quarantined from other airside locations/spaces? Who will be responsible and what levels of control will be put in place for workers entering and exiting the worksite?": "Việc kiểm soát truy cập vào công trường sẽ được duy trì như thế nào và công trường sẽ được cách ly ra sao khỏi các khu vực/không gian airside khác? Ai sẽ chịu trách nhiệm và những mức độ kiểm soát nào sẽ được áp dụng đối với công nhân ra vào công trường?",
    "What access controls will be required at remote locations (design offices, project offices within supplier organizations, offsite manufacturing facilities)?": "Những biện pháp kiểm soát truy cập nào sẽ cần có tại các địa điểm từ xa (văn phòng thiết kế, văn phòng dự án trong tổ chức của nhà cung cấp, cơ sở sản xuất ngoài công trường)?",
    "Restrictions on where (geographically) the CDE can be hosted. Preferably at the airport client’s own premises, if not then definitely within the UK.": "Các hạn chế về địa điểm đặt CDE theo phương diện địa lý. Tốt nhất là đặt tại chính cơ sở của khách hàng sân bay; nếu không thì chắc chắn phải ở trong phạm vi Vương quốc Anh.",
    "Technical:": "Công nghệ:",
    "Levels of password strength required from CDE users.": "Mức độ mạnh của mật khẩu được yêu cầu đối với người dùng CDE.",
    "Lock-down of features on client-side and supply-side hardware (laptops, tablets, phones, PCs). USB ports disabled.": "Khóa chặt các tính năng trên thiết bị phần cứng phía khách hàng và phía nhà cung cấp (laptop, máy tính bảng, điện thoại, PC). Vô hiệu hóa các cổng USB.",
    "Levels of security certification and accreditation needed for all software being used on the project – design, collaboration, project management, financial control.": "Mức độ chứng nhận và công nhận an ninh cần có đối với toàn bộ phần mềm sử dụng trong dự án - thiết kế, cộng tác, quản lý dự án, kiểm soát tài chính.",
    "Information:": "Thông tin:",
    "Description of the security classifications for all documents and information and the working practices associated with different classifications.": "Mô tả các mức phân loại an ninh áp dụng cho toàn bộ tài liệu và thông tin, cùng các thực hành làm việc tương ứng với từng mức phân loại.",
    "Description of how different types of information container in the CDE are categorized according to the security classification.": "Mô tả cách các loại container thông tin khác nhau trong CDE được phân loại theo mức phân loại an ninh.",
    "Activity 4: Responses to a security breach": "Hoạt động 4: Ứng phó với vi phạm an ninh",
    "The general list of processes to include in a security breach/incident management plan are:": "Danh sách tổng quát các quy trình cần có trong kế hoạch quản lý vi phạm/sự cố an ninh gồm:",
    "Who to contact immediately on discovery of the breach/incident and how to contact them (their details)": "Cần liên hệ ngay với ai khi phát hiện vi phạm/sự cố và liên hệ bằng cách nào (thông tin liên hệ của họ)",
    "How to work out what has been lost, damaged, or compromised": "Cách xác định những gì đã bị mất, hư hỏng hoặc bị xâm phạm",
    "How to work out who any concerned parties are (such as individuals whose personal data may have been compromised, or individuals who may have been put at more risk because of the incident)": "Cách xác định những bên nào bị ảnh hưởng hoặc có liên quan (chẳng hạn những cá nhân có dữ liệu cá nhân có thể đã bị xâm phạm, hoặc những người có thể phải đối mặt với rủi ro cao hơn do sự cố)",
    "How to liaise with any relevant third parties (such as a regulator, the media)": "Cách phối hợp với các bên thứ ba liên quan (chẳng hạn cơ quan quản lý, truyền thông)",
    "How to liaise with the general public (if it is decided that this is appropriate)": "Cách trao đổi với công chúng, nếu xác định rằng việc đó là phù hợp",
    "The project scenario for this activity is from Activity 3 – a new-build terminal at an international airport.": "Kịch bản dự án cho hoạt động này được lấy từ Hoạt động 3 - một nhà ga mới xây tại một sân bay quốc tế.",
    "The following security breach has occurred: An unauthorized person gained access to the site by using a legitimate worker’s security pass (his bag was stolen in a bar and the loss of the pass was not noticed until 2 days later as he was on leave).": "Đã xảy ra vi phạm an ninh sau đây: Một người không được phép đã vào được công trường bằng cách sử dụng thẻ an ninh hợp lệ của một công nhân (túi của người này bị lấy cắp tại một quán bar và việc mất thẻ chỉ được phát hiện sau 2 ngày vì anh ta đang nghỉ phép).",
    "The incident management plan is broken down into the following sections:": "Kế hoạch quản lý sự cố được chia thành các phần sau:",
    "1. Immediate actions": "1. Hành động ngay lập tức",
    "Notification of the loss of the pass to Head of Project Security": "Thông báo việc mất thẻ cho Trưởng bộ phận An ninh Dự án",
    "Head of Project Security immediately cancels or requests cancellation of the pass – this means the turnstiles at the site entrance will not recognize the pass.": "Trưởng bộ phận An ninh Dự án lập tức hủy hoặc yêu cầu hủy thẻ này - điều đó có nghĩa là các cổng xoay tại lối vào công trường sẽ không chấp nhận thẻ nữa.",
    "Head of Project Security asks for a report on whether and when the pass has been used since its theft – this is obtained from analysis of the site entry control system.": "Trưởng bộ phận An ninh Dự án yêu cầu báo cáo về việc thẻ đã được sử dụng hay chưa và được sử dụng khi nào kể từ lúc bị lấy cắp - thông tin này được rút ra từ việc phân tích hệ thống kiểm soát ra vào công trường.",
    "If it seems that a crime may have been committed, then the Police are informed.": "Nếu có dấu hiệu cho thấy một hành vi phạm tội có thể đã xảy ra, thì cần thông báo cho cảnh sát.",
    "2. Short term actions": "2. Hành động ngắn hạn",
    "Head of Project Security checks with Project Manager how quickly the employee needs a new pass.": "Trưởng bộ phận An ninh Dự án trao đổi với Quản lý Dự án để xác định nhân viên đó cần thẻ mới gấp đến mức nào.",
    "Head of Project Security interviews worker who lost the pass.": "Trưởng bộ phận An ninh Dự án phỏng vấn công nhân đã làm mất thẻ.",
    "Report comes back from access control with times of entry after the pass was lost. CCTV footage from the site around these times is examined for unrecognized individuals.": "Bộ phận kiểm soát ra vào cung cấp báo cáo về các thời điểm thẻ đã được dùng sau khi bị mất. Hình ảnh CCTV của công trường vào các thời điểm đó được kiểm tra để phát hiện những cá nhân không được nhận diện.",
    "The areas where the intruder has been seen are examined to check for thefts of equipment, materials, or other malicious acts.": "Các khu vực mà kẻ xâm nhập được nhìn thấy sẽ được kiểm tra để xem có xảy ra việc trộm cắp thiết bị, vật tư hay các hành vi ác ý khác hay không.",
    "If any interference or tampering with information systems is found (use of computer terminals, placing or removal of equipment) then a detailed audit of those information systems is carried out.": "Nếu phát hiện bất kỳ sự can thiệp hoặc sửa đổi trái phép nào đối với các hệ thống thông tin (sử dụng terminal máy tính, lắp đặt hoặc tháo dỡ thiết bị) thì phải tiến hành kiểm tra chi tiết các hệ thống thông tin đó.",
    "3. Medium term actions": "3. Hành động trung hạn",
    "The need for enhanced training about keeping passes secure is considered.": "Xem xét nhu cầu tăng cường đào tạo về việc giữ an toàn cho thẻ ra vào.",
    "The need for additional access control checks at key points of the site is considered.": "Xem xét nhu cầu bổ sung các bước kiểm tra kiểm soát ra vào tại những điểm trọng yếu của công trường.",
    "The integration of the access control and the staff leave records is considered.": "Xem xét việc tích hợp hệ thống kiểm soát ra vào với hồ sơ nghỉ phép của nhân viên.",
    "Activity 5: Action planning for BIM security mindedness": "Hoạt động 5: Lập kế hoạch hành động cho tư duy an ninh của BIM",
    "Focus area": "Phạm vi xem xét",
    "To our asset/information": "Đối với tài sản/thông tin của chúng ta",
    "To neighbouring assets": "Đối với tài sản lân cận",
    "Physical": "Vật lý",
    "Armed robbery\nTerrorist attack\nPolitical demonstrations against the bank’s country of origin\nPublic visitors entering unauthorized areas\nContractors entering unauthorized areas.\nStaff entering unauthorized areas.\nDisabling CCTV systems": "Cướp có vũ trang\nTấn công khủng bố\nBiểu tình chính trị chống lại quốc gia xuất xứ của ngân hàng\nKhách bên ngoài đi vào khu vực không được phép\nNhà thầu đi vào khu vực không được phép\nNhân viên đi vào khu vực không được phép\nVô hiệu hóa hệ thống CCTV",
    "Our bank building compromises boundary security for substation": "Tòa nhà ngân hàng của chúng ta làm suy yếu an ninh ranh giới của trạm biến áp",
    "Cyber": "Mạng",
    "Hacking the CAFM system to disrupt/shut down key systems.\nObtaining 3-d model O&M information to spot security weaknesses": "Xâm nhập hệ thống CAFM để gây gián đoạn hoặc tắt các hệ thống trọng yếu.\nThu thập thông tin O&M từ mô hình 3D để nhận diện các điểm yếu an ninh.",
    "Obtaining plans showing location of main cables entering substation (detail left over from site studies during construction)": "Thu thập bản vẽ thể hiện vị trí các tuyến cáp chính đi vào trạm biến áp (chi tiết còn sót lại từ các khảo sát hiện trường trong giai đoạn thi công)",
    "Area of \nISO 19650-5": "Phạm vi của\nISO 19650-5",
    "How well prepared? \n(1-10)": "Mức độ sẵn sàng đến đâu?\n(1-10)",
    "What more do we need to do?": "Chúng ta còn cần làm gì thêm?",
    "How will we know it has been done?": "Làm sao biết việc này đã hoàn thành?",
    "What resources are needed to achieve this?": "Cần những nguồn lực nào để đạt được điều này?",
    "Understanding security threats to assets and information": "Hiểu các mối đe dọa an ninh đối với tài sản và thông tin",
    "Create greater understanding amongst middle managers and project designers": "Nâng cao hiểu biết trong đội ngũ quản lý cấp trung và các nhà thiết kế dự án",
    "Include in annual skills/experience audits": "Đưa vào hoạt động rà soát kỹ năng/kinh nghiệm hằng năm",
    "Send key staff on training courses and then run in-house workshops to spread the learning": "Cử nhân sự chủ chốt tham gia các khóa đào tạo, sau đó tổ chức hội thảo nội bộ để lan tỏa kiến thức",
    "Implementing a security triage process": "Triển khai quy trình phân loại an ninh",
    "Not much as this is client task, but we should understand it": "Không nhiều vì đây là nhiệm vụ của khách hàng, nhưng chúng ta cần hiểu quy trình này",
    "Outcome from a ISO19650-5 project": "Kết quả đầu ra từ một dự án theo ISO 19650-5",
    "Read ISO19650-5 Clause 4.7 and discuss at in-house CPD": "Đọc Điều 4.7 của ISO 19650-5 và thảo luận trong buổi CPD nội bộ",
    "Identifying sensitive information and processes to handle it": "Xác định thông tin nhạy cảm và quy trình xử lý thông tin đó",
    "Create deep understanding of the types of asset listed in ISO19650-5 Clause 4.3": "Xây dựng hiểu biết sâu về các loại tài sản được liệt kê tại Điều 4.3 của ISO 19650-5",
    "Success in bidding for high-security projects": "Thành công khi đấu thầu các dự án yêu cầu an ninh cao",
    "Specialist security training for projects/buildings/ infrastructure": "Đào tạo an ninh chuyên sâu cho dự án/tòa nhà/hạ tầng",
    "Drawing up or responding to a security strategy ": "Xây dựng hoặc phản hồi đối với chiến lược an ninh",
    "3 (responding to the strategy)": "3 (phản hồi đối với chiến lược)",
    "Prepare in-house document setting out security capabilities for project information": "Chuẩn bị tài liệu nội bộ nêu rõ năng lực an ninh đối với thông tin dự án",
    "Capability statement issued to work-winning team": "Ban hành bản mô tả năng lực cho nhóm phát triển kinh doanh/đấu thầu",
    "Time (3-5 person days from business development Director and BIM champion). Some external consultant expertise to peer-review": "Thời gian (3-5 ngày công của Giám đốc Phát triển Kinh doanh và đầu mối thúc đẩy BIM). Cần thêm một phần chuyên môn tư vấn bên ngoài để phản biện.",
    "Writing or responding to a security management plan ": "Soạn thảo hoặc phản hồi đối với kế hoạch quản lý an ninh",
    "3 (Responding to the plan)": "3 (phản hồi đối với kế hoạch)",
    "Part of in-house security capability statement": "Là một phần của bản mô tả năng lực an ninh nội bộ",
    "See above": "Xem ở trên",
    "Implementing security requirements in the CDE": "Triển khai các yêu cầu an ninh trong CDE",
    "Make sure all project staff understand client’s security designations and procedures for each one": "Bảo đảm toàn bộ nhân sự dự án hiểu rõ các mức phân loại an ninh của khách hàng và quy trình tương ứng cho từng mức",
    "In house audits on live projects": "Kiểm tra nội bộ trên các dự án đang triển khai",
    "Training for designers, information managers": "Đào tạo cho nhà thiết kế và người quản lý thông tin",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Translate Part 5 references DOCX to Vietnamese.")
    parser.add_argument("--input", required=True, help="Input DOCX path.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    return parser.parse_args()


def normalize(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", "").strip()


EXACT_TRANSLATIONS = {normalize(key): value for key, value in RAW_EXACT_TRANSLATIONS.items()}


def cleanup_translation(text: str) -> str:
    replacements = (
        ("bảo mật", "an ninh"),
        ("tư duy bảo mật", "tư duy an ninh"),
        ("Thực hành thực tế", "Triển khai"),
        ("an ninh thẻ", "thẻ an ninh"),
        ("khách công cộng", "khách bên ngoài"),
        ("vi phạm/sự cố bảo mật", "vi phạm/sự cố an ninh"),
        ("kế hoạch quản lý bảo mật", "kế hoạch quản lý an ninh"),
        ("người dùng CDE", "người dùng CDE"),
    )
    cleaned = text
    for old, new in replacements:
        cleaned = cleaned.replace(old, new)
    return cleaned


def translate_text(text: str, translator: GoogleTranslator, cache: dict[str, str]) -> tuple[str, str]:
    key = normalize(text)
    if not key:
        return text, "empty"

    if key in EXACT_TRANSLATIONS:
        return EXACT_TRANSLATIONS[key], "exact"

    if key in cache:
        return cache[key], "cache"

    translated = translator.translate(key) or key
    translated = cleanup_translation(translated)
    cache[key] = translated
    return translated, "fallback"


def main() -> None:
    args = parse_args()
    source = Path(args.input)
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    copyfile(source, output)

    doc = Document(output)
    translator = GoogleTranslator(source="en", target="vi")
    cache: dict[str, str] = {}
    stats = {"exact": 0, "fallback": 0, "cache": 0, "empty": 0}

    for paragraph in doc.paragraphs:
        original = paragraph.text
        translated, mode = translate_text(original, translator, cache)
        stats[mode] += 1
        if normalize(original):
            paragraph.text = translated

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text
                translated, mode = translate_text(original, translator, cache)
                stats[mode] += 1
                if normalize(original):
                    cell.text = translated

    doc.save(output)
    print(f"Translated DOCX saved to: {output}")
    print(stats)


if __name__ == "__main__":
    main()
